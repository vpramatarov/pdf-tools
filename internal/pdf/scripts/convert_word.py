import sys
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
import re

def remove_control_characters(text):
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)


def pdf_to_linear_docx(pdf_path, docx_path, sort_mode):
    doc_word = Document()
    
    try:
        pdf_document = fitz.open(pdf_path)
    except Exception as e:
        print(f"Error opening PDF: {e}", file=sys.stderr)
        sys.exit(1)

    body_buffer = []

    def flush_buffer():
        if not body_buffer: return
        raw_text = "\n".join(body_buffer)
        hyphen_pattern = r'(?<=[а-яА-Яa-zA-Z])-\s*\n\s*(?=[а-яА-Яa-zA-Z])'
        text_dehyphenated = re.sub(hyphen_pattern, '', raw_text)
        linear_text = text_dehyphenated.replace('\n', ' ')
        clean_text = re.sub(r'\s+', ' ', linear_text).strip()
        if not clean_text:
            body_buffer.clear()
            return
        split_pattern = r'(?<=[.!?])\s+(?=[А-ЯA-Z])'
        paragraphs = re.split(split_pattern, clean_text)
        for para_text in paragraphs:
            if not para_text.strip(): continue
            p = doc_word.add_paragraph()
            runner = p.add_run(para_text.strip())
            runner.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)
        body_buffer.clear()

    for page_num, page in enumerate(pdf_document):
        blocks = page.get_text("blocks", sort=sort_mode)

        for block in blocks:
            text = block[4].strip()
            text = remove_control_characters(text)
            
            if not text: continue
            
            text_lower = text.lower().strip()
            if ("на стр." in text_lower) or ("от стр." in text_lower): 
                continue

            if (text_lower == "квантов" or text_lower == "преход") or ("квантов" in text_lower or "преход" in text_lower):
                continue

            if text_lower.isdigit():
                continue

            is_all_caps = text.isupper() and len(text) > 2
            ends_with_sentence_punct = text.endswith(('.', ',', ':', ';', '!', '?'))
            ends_with_hyphen = text.endswith('-')
            is_title_candidate = (len(text) < 150 and not ends_with_sentence_punct and not ends_with_hyphen)
            
            if is_all_caps or is_title_candidate:
                flush_buffer()
                clean_title = text.replace('\n', ' ').replace('- ', '')
                clean_title = re.sub(r'\s+', ' ', clean_title).strip()
                clean_title = remove_control_characters(clean_title)

                if clean_title.lower() in ["квантов", "преход", "източник:"]: 
                    continue
                
                p = doc_word.add_paragraph()
                runner = p.add_run(clean_title)
                runner.bold = True
                runner.font.size = Pt(12)
                p.paragraph_format.space_after = Pt(12)
            else:
                body_buffer.append(text)

    flush_buffer()

    try:
        doc_word.save(docx_path)
        print(f"Successfully saved to {docx_path}")
    except Exception as e:
        print(f"Error saving DOCX: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    # Expect 3 arguments (script + input + output + sort)
    # If the 4th argument (index 3) is not provided, default to True
    if len(sys.argv) < 3:
        print("Usage: python convert_word.py <input_pdf> <output_docx> [sort_mode]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    # Read the sort argument
    sort_arg = True
    if len(sys.argv) >= 4:
        # Accept 'false', '0', 'no' as False, everything else is True
        sort_arg = sys.argv[3].lower() not in ('false', '0', 'no')

    pdf_to_linear_docx(input_file, output_file, sort_arg)